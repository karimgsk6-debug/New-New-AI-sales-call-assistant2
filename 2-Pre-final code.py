# app.py - Final merged app with inline citation snippets (TF-IDF), Prompt suggestions,
# background + header logos, multi-brand support, PDF upload, TTS, exports, and improved UI.

import streamlit as st
from PIL import Image
import re, os, tempfile, base64, requests, math
from io import BytesIO
from datetime import datetime
from groq import Groq
from PyPDF2 import PdfReader
from html import escape
from gtts import gTTS

# TF-IDF tools
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import linear_kernel

# Optional DOCX export
try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# ElevenLabs fallback flag
try:
    import elevenlabs
    ELEVENLABS_AVAILABLE = True
except Exception:
    ELEVENLABS_AVAILABLE = False

# ---------------------------- Page config ----------------------------
st.set_page_config(page_title="AI Sales Call Assistant", layout="wide")

# ---------- Repo info for linking files ----------
REPO_USER = "karimgsk6-debug"
REPO_NAME = "New-New-AI-sales-call-assistant2"
COMMIT = "845b8f1ae98e46440e840c0a906f3610dd343c9a"
REPO_BLOB_BASE = f"https://github.com/{REPO_USER}/{REPO_NAME}/blob/{COMMIT}/.devcontainer"
REPO_RAW_BASE = f"https://raw.githubusercontent.com/{REPO_USER}/{REPO_NAME}/{COMMIT}/.devcontainer"

# ---------------------------- Assets (raw URLs) ----------------------------
# Note: this path worked in prior messages; if it doesn't resolve for you, place images in repo and use local st.image()
BACKGROUND_URL = REPO_RAW_BASE + "/.devcontainer/background1.png"
GSK_LOGO_RAW = f"https://raw.githubusercontent.com/{REPO_USER}/{REPO_NAME}/main/.devcontainer/GSK1-logo.png"
AI_LOGO_RAW = f"https://raw.githubusercontent.com/{REPO_USER}/{REPO_NAME}/main/.devcontainer/AURA1.png"

# ---------------------------- Session defaults ----------------------------
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "uploaded_pdf_text" not in st.session_state:
    st.session_state.uploaded_pdf_text = ""
if "pdf_summary" not in st.session_state:
    st.session_state.pdf_summary = ""
if "voice_pref" not in st.session_state:
    st.session_state.voice_pref = "Old Male"
if "pdf_summary_size" not in st.session_state:
    st.session_state.pdf_summary_size = "Normal"
if "main_input" not in st.session_state:
    st.session_state.main_input = ""
if "selected_brand" not in st.session_state:
    st.session_state.selected_brand = "trelegy"

# ---------------------------- CSS / layout ----------------------------
CSS = f"""
<style>
/* App background */
.stApp {{
  background-image: url('{BACKGROUND_URL}');
  background-size: cover;
  background-attachment: fixed;
  background-position: center;
}}

/* Title box */
.title-box {{
  background: rgba(255,255,255,0.92);
  padding: 12px;
  border-radius: 10px;
  margin-bottom: 14px;
  position: relative;
  display:flex;
  align-items:center;
  justify-content:center;
}}
.title-box img.left-logo {{ position:absolute; left:16px; width:130px; height:auto; }}
.title-box img.right-logo {{ position:absolute; right:16px; width:130px; height:auto; }}
.title-box h1 {{ margin:0; font-size:22px; }}

/* Main chat area */
.chat-container {{
  max-height: 56vh;
  overflow-y:auto;
  padding: 14px;
  border-radius: 10px;
  background: rgba(255,255,255,0.94);
  margin-bottom: 140px;
}}
.chat-bubble-user {{ background:#0078D7; color:white; padding:12px; border-radius:10px; margin:8px 0; max-width:78%; margin-left:auto; }}
.chat-bubble-ai {{ background:#d9f0ff; color:#000; padding:12px; border-radius:10px; margin:8px 0; max-width:78%; }}

/* Suggestions area (collapsed by Streamlit expander, but style inner area) */
.suggestions-inline {{ background: rgba(255,255,255,0.96); padding:10px; border-radius:8px; box-shadow:0 4px 10px rgba(0,0,0,0.06); }}
.suggestion-pill {{ background:#fff; border:1px solid #ddd; padding:8px 12px; border-radius:20px; cursor:pointer; margin:6px; display:inline-block; }}
.suggestion-pill:hover {{ background:#eef6ff; }}

/* Input area fixed at bottom */
.input-area {{
  position: fixed;
  left:24px;
  right:24px;
  bottom:18px;
  z-index:9999;
  display:flex;
  gap:8px;
  align-items:flex-end;
}}
.input-area textarea {{
  width:100%;
  min-height:72px;
  max-height:200px;
  padding:10px;
  border-radius:8px;
  border:1px solid #ccc;
  resize:vertical;
}}
.send-button {{
  height:44px;
  padding:0 14px;
  border-radius:8px;
  border:none;
  background:#FF6F00;
  color:white;
  cursor:pointer;
  display:flex;
  align-items:center;
  gap:8px;
  font-weight:600;
}}

/* Inline citation styling */
.citation-box {{
  background:#fbfbff;
  border-left:4px solid #0078D7;
  padding:8px;
  margin-top:8px;
  border-radius:6px;
  font-size:13px;
  white-space:pre-wrap;
}}

/* disclaimer */
.fixed-disclaimer {{ position:fixed; left:0; right:0; bottom:0; background:rgba(255,255,255,0.95); padding:8px; border-top:2px solid #FF6F00; text-align:center; font-size:12px; z-index:9997; }}
</style>
"""
st.markdown(CSS, unsafe_allow_html=True)

# ---------------------------- GROQ client init (safe) ----------------------------
GROQ_API_KEY = st.secrets.get("GROQ_API_KEY", "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")
client = None
if GROQ_API_KEY:
    try:
        client = Groq(api_key=GROQ_API_KEY)
    except Exception:
        client = None

# ---------------------------- Brands ----------------------------
brand_data = {
    "trelegy": {
        "display": "Trelegy",
        "segments": ["Awareness", "Diagnosis", "Adoption", "Adherence"],
        "personas": ["Primary Care COPD Prescriber", "Pulmonologist", "Respiratory Nurse"],
        "barriers": ["Formulary access", "Inhaler technique", "Concerns about side effects", "Cost/coverage"],
        "references_path": ".devcontainer/references/trelegy/",
        "sales_path": ".devcontainer/SalesModule/trelegy"
    },
    "shingrix": {
        "display": "Shingrix",
        "segments": ["R – Reach", "A – Acquisition", "C – Conversion", "E – Engagement"],
        "personas": ["Uncommitted Vaccinator", "Reluctant Efficiency", "Patient Influenced", "Committed Vaccinator"],
        "barriers": ["HCP does not consider HZ a risk", "No time for discussion", "Cost concerns", "Not convinced of efficacy"],
        "references_path": ".devcontainer/references/shingrix/",
        "sales_path": ".devcontainer/SalesModule/shingrix"
    },
    "jemperli": {
        "display": "Jemperli",
        "segments": ["Target Identification", "Trial Adoption", "Routine Use", "Advocacy"],
        "personas": ["Data-Driven Oncologist", "Skeptical Specialist", "Innovator Prescriber", "Late Adopter"],
        "barriers": ["Unfamiliar with immunotherapy", "Safety concerns", "Limited patient eligibility", "Access/reimbursement issues"],
        "references_path": ".devcontainer/references/jemperli/",
        "sales_path": ".devcontainer/SalesModule/jemperli"
    }
}

specialties = ["GP", "Pulmonologist", "Internal medicine", "Oncologist"]
objectives = ["Awareness", "Adoption", "Retention"]

# ---------------------------- Helper functions for local docs & TF-IDF ----------------------------
def read_file_text(path):
    try:
        if path.lower().endswith(".pdf"):
            reader = PdfReader(path)
            text = "".join([p.extract_text() or "" for p in reader.pages])
        else:
            with open(path, "r", encoding="utf-8", errors="ignore") as fh:
                text = fh.read()
        return text
    except Exception as e:
        return f"[Error reading {os.path.basename(path)}: {e}]"

def build_corpus_for_paths(folder_paths, chunk_size_sentences=3):
    """
    Reads all txt/pdf files from folder_paths (list) and returns:
      - chunks: list of text chunks (strings)
      - metadatas: list of dicts {filename, start_sentence_index}
    """
    chunks = []
    metadatas = []
    for folder in folder_paths:
        if not folder or not os.path.exists(folder):
            continue
        files = [f for f in os.listdir(folder) if f.lower().endswith((".pdf", ".txt"))]
        for fname in files:
            p = os.path.join(folder, fname)
            text = read_file_text(p)
            # split into sentences (simple)
            sents = re.split(r'(?<=[\.\?\!])\s+', text)
            # create sliding windows of chunk_size_sentences
            for i in range(0, max(1, len(sents)), chunk_size_sentences):
                chunk = " ".join(sents[i:i+chunk_size_sentences]).strip()
                if chunk:
                    chunks.append(chunk)
                    metadatas.append({"filename": fname, "folder": folder, "start": i})
    return chunks, metadatas

def find_top_n_snippets(query, chunks, metadatas, top_n=3):
    if not chunks:
        return []
    try:
        vectorizer = TfidfVectorizer(stop_words="english").fit(chunks + [query])
        chunk_vecs = vectorizer.transform(chunks)
        q_vec = vectorizer.transform([query])
        # compute cosine similarities
        sims = linear_kernel(q_vec, chunk_vecs).flatten()
        top_idxs = sims.argsort()[::-1][:top_n]
        results = []
        for idx in top_idxs:
            if sims[idx] <= 0:
                continue
            results.append({"score": float(sims[idx]), "text": chunks[idx], "meta": metadatas[idx]})
        return results
    except Exception:
        # fallback simple substring matching
        out = []
        q = query.lower()
        for i, c in enumerate(chunks):
            if q in c.lower():
                out.append({"score": 1.0, "text": c, "meta": metadatas[i]})
                if len(out) >= top_n: break
        return out

# ---------------------------- Audio generation helper ----------------------------
ELEVENLABS_API_KEY = st.secrets.get("ELEVENLABS_API_KEY", "")
ELEVENLABS_VOICE_ID = st.secrets.get("ELEVENLABS_VOICE_ID", "")
if ELEVENLABS_AVAILABLE and ELEVENLABS_API_KEY and ELEVENLABS_VOICE_ID:
    try:
        elevenlabs.api_key = ELEVENLABS_API_KEY
    except:
        pass

def generate_audio(text):
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3")
    try:
        tts_text = re.sub(r'[,*]{1,}', '', text)
        if ELEVENLABS_AVAILABLE and ELEVENLABS_API_KEY and ELEVENLABS_VOICE_ID:
            audio_stream = elevenlabs.generate(text=tts_text, voice=ELEVENLABS_VOICE_ID, stream=True)
            with open(tmp.name, "wb") as f:
                for ch in audio_stream:
                    f.write(ch)
        else:
            tts = gTTS(text=tts_text, lang="en", slow=False)
            tts.save(tmp.name)
        with open(tmp.name, "rb") as fh:
            return base64.b64encode(fh.read()).decode()
    except Exception as e:
        st.warning(f"Audio generation failed: {e}")
        return ""

# ---------------------------- Sidebar: filters, clear chat, language ----------------------------
with st.sidebar.expander("Filters & Options", expanded=True):
    sel_brand_key = st.selectbox("Brand", sorted(list(brand_data.keys())),
                                index=list(sorted(brand_data.keys())).index(st.session_state.get("selected_brand","trelegy")))
    st.session_state.selected_brand = sel_brand_key
    sel_brand = brand_data[sel_brand_key]
    segment = st.selectbox("Segment", sel_brand["segments"], key="sidebar_segment")
    persona = st.selectbox("HCP Persona", sel_brand["personas"], key="sidebar_persona")
    barrier = st.multiselect("Doctor Barrier", sel_brand["barriers"], key="sidebar_barrier")
    specialty = st.selectbox("Specialty", specialties, key="sidebar_specialty")
    objective = st.selectbox("Objective", objectives, key="sidebar_objective")
    response_tone = st.selectbox("Response Tone", ["Formal", "Casual", "Friendly", "Persuasive"], key="sidebar_tone")
    response_length = st.selectbox("Response Length", ["Short", "Medium", "Long"], key="sidebar_length")
    # language widget - store value to a variable (do not set same key in session_state)
    selected_language = st.radio("Language", ["English", "Arabic"], horizontal=True, key="sidebar_language")
    if st.button("🗑️ Clear Chat", key="sidebar_clear"):
        st.session_state.chat_history = []
with st.sidebar.expander("🌐 Add External Reference URLs", expanded=False):
    external_urls = st.text_area("Enter URLs (one per line)").splitlines()
with st.sidebar.expander("📄 Export Options", expanded=False):
    export_format = st.radio("Choose Export Format", ["TXT", "DOCX"], horizontal=True)

# ---------------------------- Title / header box ----------------------------
st.markdown(f"""
<div class="title-box">
  <img src="{GSK_LOGO_RAW}" class="left-logo">
  <h1>💡 AI Sales Call Assistant — {brand_data[st.session_state.selected_brand]['display']}</h1>
  <img src="{AI_LOGO_RAW}" class="right-logo">
</div>
""", unsafe_allow_html=True)

# ---------------------------- Load local files (refs + sales) and build TF-IDF corpus ----------
refs_folder = brand_data[st.session_state.selected_brand]["references_path"]
sales_folder = brand_data[st.session_state.selected_brand]["sales_path"]

local_refs_text, local_ref_files = "", []
sales_text, sales_files = "", []
if os.path.exists(refs_folder):
    local_refs_text, local_ref_files = "", []
    # load combined text and list of files
    for f in os.listdir(refs_folder):
        if f.lower().endswith((".pdf", ".txt")):
            local_ref_files.append(f)
            try:
                local_refs_text += read_file_text(os.path.join(refs_folder, f)) + "\n"
            except:
                pass

if os.path.exists(sales_folder):
    sales_text, sales_files = "", []
    for f in os.listdir(sales_folder):
        if f.lower().endswith((".pdf", ".txt")):
            sales_files.append(f)
            try:
                sales_text += read_file_text(os.path.join(sales_folder, f)) + "\n"
            except:
                pass

# Build corpus chunks once per run (cheap)
corpus_folders = []
if refs_folder: corpus_folders.append(refs_folder)
if sales_folder: corpus_folders.append(sales_folder)
chunks, chunk_meta = build_corpus_for_paths(corpus_folders, chunk_size_sentences=3)

# ---------------------------- PDF Upload area ----------------------------
with st.expander("📄 Upload Custom PDF for AI Context", expanded=False):
    uploaded_pdf = st.file_uploader("Upload PDF", type=["pdf"])
    st.session_state.pdf_summary_size = st.radio("PDF Summary Size", ["Consisted", "Normal", "Detailed"],
                                                horizontal=True, key="pdf_size_widget")
    if uploaded_pdf:
        try:
            reader = PdfReader(uploaded_pdf)
            full_text = "".join([p.extract_text() or "" for p in reader.pages])
            st.session_state.uploaded_pdf_text = full_text
            st.success(f"Loaded {len(full_text)} characters from uploaded PDF.")
            bullets_count = {"Consisted":5,"Normal":10,"Detailed":20}.get(st.session_state.pdf_summary_size,10)
            # try to use model for summarization if present
            if client:
                try:
                    summ_prompt = f"Summarize into {bullets_count} bullet points:\n{full_text[:12000]}"
                    summ = client.chat.completions.create(model="llama-3.3-70b-versatile",
                                                         messages=[{"role":"user","content":summ_prompt}], temperature=0.4)
                    st.session_state.pdf_summary = summ.choices[0].message.content
                except Exception:
                    sts = re.findall(r'([A-Z][^.]{20,200})', full_text)
                    st.session_state.pdf_summary = "\n".join(sts[:bullets_count])
            else:
                sts = re.findall(r'([A-Z][^.]{20,200})', full_text)
                st.session_state.pdf_summary = "\n".join(sts[:bullets_count])
        except Exception as e:
            st.error(f"Error reading uploaded PDF: {e}")
    if st.session_state.pdf_summary:
        st.markdown(f"<div style='background:#E6F0FF;padding:12px;border-radius:8px;white-space:pre-line'>{escape(st.session_state.pdf_summary)}</div>", unsafe_allow_html=True)

# ---------------------------- Chat history display ----------------------------
st.markdown('<div class="chat-container">', unsafe_allow_html=True)
for msg in st.session_state.chat_history:
    if msg.get("role") == "user":
        st.markdown(f'<div class="chat-bubble-user">🧑 You: {escape(msg.get("content",""))}</div>', unsafe_allow_html=True)
    else:
        st.markdown(f'<div class="chat-bubble-ai">🤖 AI: {escape(msg.get("content",""))}</div>', unsafe_allow_html=True)
        # render inline citations if present
        if msg.get("citations"):
            # citations is a list of dicts {text, meta}
            for c in msg["citations"]:
                fname = c["meta"]["filename"]
                blob_url = f"{REPO_BLOB_BASE}/references/{st.session_state.selected_brand}/{fname}" if fname in local_ref_files else f"{REPO_BLOB_BASE}/SalesModule/{st.session_state.selected_brand}/{fname}"
                st.markdown(f'<div class="citation-box"><b>Excerpt from {escape(fname)}:</b><br>{escape(c["text"][:800])}...<br><a href="{blob_url}" target="_blank">View full file</a></div>', unsafe_allow_html=True)
        if msg.get("audio"):
            try:
                st.audio(base64.b64decode(msg["audio"]), format="audio/mp3")
            except:
                pass
st.markdown('</div>', unsafe_allow_html=True)

# ---------------------------- Prompt suggestions (collapsed expander above input) ------------
def build_suggestions_for_brand(brand_key, persona, barrier_list, segment, specialty, objective):
    s = []
    s.append(f"Generate call flow for {persona} focused on {objective}.")
    if barrier_list:
        s.append(f"Handle objection: {', '.join(barrier_list[:2])} for {persona}.")
    else:
        s.append(f"Identify common objections for {persona}.")
    s.append(f"Summarize HCP persona insights for {persona}.")
    s.append(f"Key talking points for {brand_data[brand_key]['display']} in {segment}.")
    s.append(f"Draft a short adoption message for {brand_data[brand_key]['display']} to a {specialty}.")
    return s

with st.expander("Prompt Suggestions (click to autofill)", expanded=False):
    suggs = build_suggestions_for_brand(st.session_state.selected_brand, persona, barrier, segment, specialty, objective)
    st.markdown('<div class="suggestions-inline">', unsafe_allow_html=True)
    # render as 3 columns for better visual layout
    cols = st.columns([1,1,1])
    for i, s in enumerate(suggs):
        col = cols[i % 3]
        if col.button(s, key=f"sugg_{i}"):
            st.session_state.main_input = s
    st.markdown('</div>', unsafe_allow_html=True)

# ---------------------------- Chat input (fixed bottom) with Send icon/button ----------------------------
with st.form(key="chat_form", clear_on_submit=False):
    # Controlled text area: populate from session_state.main_input
    message = st.text_area("Message (editable)", value=st.session_state.get("main_input",""), key="chat_input_area", height=110)
    send = st.form_submit_button("Send")
    if send:
        user_text = message.strip()
        if user_text:
            # append user message
            st.session_state.chat_history.append({"role":"user","content":user_text})
            # clear main input value safely
            st.session_state.main_input = ""
            # Build combined context
            combined_context = "\n".join([
                local_refs_text or "",
                sales_text or "",
                external_text or "",
                st.session_state.uploaded_pdf_text or ""
            ])[:15000]

            # call_flow depending on brand
            call_flow_prompt = ""
            if st.session_state.selected_brand.lower() == "jemperli":
                call_flow_prompt = "\n".join([f"{k}: {v}" for k,v in jemperli_CALL_FLOW.items()])
            elif st.session_state.selected_brand.lower() == "shingrix":
                call_flow_prompt = "\n".join([f"{k}: {v}" for k,v in shingrix_CALL_FLOW.items()])
            elif st.session_state.selected_brand.lower() == "trelegy":
                call_flow_prompt = "\n".join([f"{k}: {v}" for k,v in trelegy_CALL_FLOW.items()])

            # model prompt
            system_prompt = "You are a pharmaceutical AI assistant. Tailor responses using references, sales modules, uploaded PDFs, and follow brand call flow."
            final_prompt = f"{user_text}\n\nBrand: {st.session_state.selected_brand}\nPersona: {persona}\nSegment: {segment}\nSpecialty: {specialty}\nObjective: {objective}\nBarriers: {', '.join(barrier) if barrier else 'None'}\n\n{call_flow_prompt}\n\nContext (truncated):\n{combined_context[:5000]}"

            # call model if available
            assistant_text = ""
            if client:
                try:
                    resp = client.chat.completions.create(
                        model="llama-3.3-70b-versatile",
                        messages=[{"role":"system","content":system_prompt},
                                  {"role":"user","content":final_prompt}],
                        temperature=0.62
                    )
                    assistant_text = resp.choices[0].message.content
                except Exception as e:
                    assistant_text = f"(AI Error) {e}"
            else:
                assistant_text = f"(Fallback) Based on local refs: {user_text}"

            # Inline citation: run TF-IDF search of chunks built earlier using user_text
            top_snips = find_top_n_snippets(user_text, chunks, chunk_meta, top_n=3)
            # prepare citations simplified for storing in chat history
            citations_for_history = top_snips if top_snips else []

            # create sources footer linking to top files
            def make_links_list(files_list, repo_subpath):
                out = []
                for fname in files_list[:3]:
                    blob = f"{REPO_BLOB_BASE}/{repo_subpath}/{fname}"
                    out.append(f"- [{fname}]({blob})")
                return "\n".join(out) if out else ""

            refs_links_txt = make_links_list(local_ref_files, f"references/{st.session_state.selected_brand}")
            sales_links_txt = make_links_list(sales_files, f"SalesModule/{st.session_state.selected_brand}")
            sources_footer = "\n\n**Sources:**\n"
            if refs_links_txt:
                sources_footer += f"\n**References:**\n{refs_links_txt}"
            if sales_links_txt:
                sources_footer += f"\n\n**Sales Modules:**\n{sales_links_txt}"

            assistant_with_sources = assistant_text + sources_footer

            # generate audio (best-effort)
            audio_b64 = ""
            try:
                audio_b64 = generate_audio(assistant_text)
            except:
                audio_b64 = ""

            st.session_state.chat_history.append({
                "role":"assistant",
                "content": assistant_with_sources,
                "audio": audio_b64,
                "citations": citations_for_history
            })

# ---------------------------- Export / Download Chat ----------
if st.session_state.chat_history:
    with st.expander("💾 Export / Download Chat", expanded=False):
        text_export = "\n\n".join([f"{e['role'].capitalize()}: {e['content']}" for e in st.session_state.chat_history])
        st.download_button("⬇️ Download TXT", text_export.encode(), file_name=f"{st.session_state.selected_brand}_chat.txt")
        if DOCX_AVAILABLE:
            if st.button("Export as DOCX"):
                try:
                    doc = Document()
                    doc.add_heading("AI Sales Call Assistant Export", 0)
                    doc.add_paragraph(f"Brand: {st.session_state.selected_brand.upper()} | Date: {datetime.now().strftime('%Y-%m-%d')}")
                    for e in st.session_state.chat_history:
                        doc.add_paragraph(f"{e['role'].capitalize()}: {e['content']}")
                    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                    doc.save(tmp.name)
                    with open(tmp.name, "rb") as fh:
                        st.download_button("⬇️ Download DOCX", fh.read(), file_name=f"{st.session_state.selected_brand}_chat.docx")
                except Exception as e:
                    st.error(f"Could not export DOCX: {e}")

# ---------------------------- Call flows (kept) ----------------------------
jemperli_CALL_FLOW = {
    "COCO": "Pre-call planning using customer insights...",
    "Anchor": "Open conversation with a patient-focused narrative...",
    "Engage": "Draw customer in through two-way dialogue...",
    "Close": "Gain agreement, define next steps..."
}
shingrix_CALL_FLOW = {
    "Prepare": "Plan the call: identify persona, objectives...",
    "Engage": "Start conversation, capture attention...",
    "Create Opportunities": "Identify gaps or unmet needs...",
    "Influence": "Present evidence, handle objections..."
}
trelegy_CALL_FLOW = {
    "Prepare": "Assess inhaler technique and adherence...",
    "Engage": "Open with symptom-based questions...",
    "Demonstrate": "Discuss inhaler technique and education...",
    "Address Access": "Clarify formulary and reimbursement options...",
    "Close": "Agree on next steps..."
}

# ---------------------------- Disclaimer ----------------------------
st.markdown('<div class="fixed-disclaimer">⚠️ This AI tool is for informational purposes only. Verify all medical content with approved references before use.</div>', unsafe_allow_html=True)



# app.py - Full AI Sales Call Assistant (Enhanced)
import streamlit as st
import os, re, tempfile, base64, io
from datetime import datetime
from html import escape

# Optional libs
try:
    from groq import Groq
except:
    Groq = None

try:
    from PyPDF2 import PdfReader
except:
    PdfReader = None

try:
    from gtts import gTTS
except:
    gTTS = None

try:
    from docx import Document
    DOCX_AVAILABLE = True
except:
    DOCX_AVAILABLE = False

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import linear_kernel
    SKLEARN_AVAILABLE = True
except:
    SKLEARN_AVAILABLE = False

try:
    import elevenlabs
    ELEVENLABS_AVAILABLE = True
except:
    ELEVENLABS_AVAILABLE = False

try:
    import pyttsx3
    PYTTSX3_AVAILABLE = True
except:
    PYTTSX3_AVAILABLE = False

# -------------------------
# Page config & background
# -------------------------
st.set_page_config(page_title="AI Sales Call Assistant", layout="wide", initial_sidebar_state="expanded")

REPO_USER = "karimgsk6-debug"
REPO_NAME = "New-New-AI-sales-call-assistant2"
COMMIT = "845b8f1ae98e46440e840c0a906f3610dd343c9a"
REPO_BLOB_BASE = f"https://github.com/{REPO_USER}/{REPO_NAME}/blob/{COMMIT}/.devcontainer"
REPO_RAW_BASE = f"https://raw.githubusercontent.com/{REPO_USER}/{REPO_NAME}/{COMMIT}/.devcontainer"
BACKGROUND_URL = REPO_RAW_BASE + "/background1.png"
GSK_LOGO_RAW = f"https://raw.githubusercontent.com/{REPO_USER}/{REPO_NAME}/main/.devcontainer/GSK1-logo.png"
AI_LOGO_RAW = f"https://raw.githubusercontent.com/{REPO_USER}/{REPO_NAME}/main/.devcontainer/AURA1.png"

# -------------------------
# Session defaults
# -------------------------
defaults = {
    "chat_history": [],
    "main_input": "",
    "selected_brand": "shingrix",
    "temperature": 0.62,
    "search_mode": "deep",
    "medical_summary": "",
    "sales_summary": "",
    "uploaded_pdf_text": "",
    "pdf_summary": "",
    "feedback": {},
    "language": "English",
}
for k,v in defaults.items():
    st.session_state.setdefault(k,v)

# -------------------------
# CSS & background
# -------------------------
CSS = f"""
<style>
[data-testid="stAppViewContainer"] {{
  background-image: url('{BACKGROUND_URL}');
  background-size: cover;
  background-position: center;
  background-attachment: fixed;
}}
.title-box {{
  background: rgba(255,255,255,0.95);
  padding: 12px;
  border-radius: 10px;
  margin-bottom: 12px;
  position: relative;
  display:flex;
  align-items:center;
  justify-content:center;
}}
.title-box img.left-logo {{ position:absolute; left:12px; height:64px; }}
.title-box img.right-logo {{ position:absolute; right:12px; height:64px; }}
.chat-container {{ max-height: 60vh; overflow-y:auto; padding:12px; background: rgba(255,255,255,0.95); border-radius:8px; margin-bottom:160px; }}
.chat-bubble-user {{ background:#0078D7; color:white; padding:10px; border-radius:12px; margin:8px 0; max-width:78%; margin-left:auto; }}
.chat-bubble-ai {{ background:#eef9ff; color:#000; padding:10px; border-radius:12px; margin:8px 0; max-width:78%; }}
.suggestion-pill {{ background:#fff; border:1px solid #ddd; padding:8px 12px; border-radius:20px; margin:6px; cursor:pointer; display:inline-block; }}
.suggestion-pill:hover {{ background:#f0f8ff; }}
.citation-box {{ background:#fbfbff; border-left:4px solid #0078D7; padding:8px; margin-top:8px; border-radius:6px; font-size:13px; white-space:pre-wrap; }}
.input-area {{ position: fixed; left:20px; right:20px; bottom:18px; z-index:9999; display:flex; gap:8px; align-items:flex-end; }}
.input-area textarea {{ width:100%; min-height:72px; max-height:250px; padding:10px; border-radius:8px; border:1px solid #ccc; resize:vertical; }}
.send-button {{ height:44px; padding:0 14px; border-radius:8px; border:none; background:#FF6F00; color:white; cursor:pointer; font-weight:600; }}
.feedback-buttons button {{ margin-right:6px; }}
.fixed-disclaimer {{ position:fixed; left:0; right:0; bottom:0; background:rgba(255,255,255,0.95); padding:8px; border-top:2px solid #FF6F00; text-align:center; font-size:12px; z-index:9997; }}
</style>
"""
st.markdown(CSS, unsafe_allow_html=True)

# -------------------------
# Initialize GROQ client
# -------------------------
GROQ_API_KEY = "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"
client = None
if Groq and GROQ_API_KEY:
    try:
        client = Groq(api_key=GROQ_API_KEY)
    except:
        client = None

# -------------------------
# Brand info
# -------------------------
brand_data = {
    "shingrix": {
        "display":"Shingrix",
        "segments":["R – Reach","A – Acquisition","C – Conversion","E – Engagement"],
        "personas":["Uncommitted Vaccinator","Reluctant Efficiency","Patient Influenced","Committed Vaccinator"],
        "barriers":["HCP does not consider HZ a risk","No time for discussion","Cost concerns","Not convinced of efficacy"],
        "specialties":["GP","Dermatologist","Geriatrician"],
        "references_path":".devcontainer/references/shingrix/",
        "sales_path":".devcontainer/SalesModule/shingrix/",
        "call_flow":["Prepare","Engage","Create Opportunities","Influence","Impact GSO","Post-call Analysis"]
    },
    "jemperli": {
        "display":"Jemperli",
        "segments":["Target Identification","Trial Adoption","Routine Use","Advocacy"],
        "personas":["Data-Driven Oncologist","Skeptical Specialist","Innovator Prescriber","Late Adopter"],
        "barriers":["Unfamiliar with immunotherapy","Safety concerns","Limited eligibility","Access/reimbursement issues"],
        "specialties":["Oncologist","Medical Oncologist"],
        "references_path":".devcontainer/references/jemperli/",
        "sales_path":".devcontainer/SalesModule/jemperli/",
        "call_flow":["COCO","Anchor","Engage","Close"]
    },
    "trelegy": {
        "display":"Trelegy",
        "segments":["Awareness","Diagnosis","Adoption","Adherence"],
        "personas":["Primary Care COPD Prescriber","Pulmonologist","Respiratory Nurse"],
        "barriers":["Formulary access","Inhaler technique","Side effect concerns","Cost/coverage"],
        "specialties":["GP","Pulmonologist","Respiratory Specialist"],
        "references_path":".devcontainer/references/trelegy/",
        "sales_path":".devcontainer/SalesModule/trelegy/",
        "call_flow":["Prepare","Engage","Demonstrate","Address Access","Close"]
    }
}

# -------------------------
# Helper functions
# -------------------------
def read_file_text(path):
    if not os.path.exists(path): return ""
    try:
        if path.lower().endswith(".pdf") and PdfReader:
            reader = PdfReader(path)
            return "".join([p.extract_text() or "" for p in reader.pages])
        else:
            with open(path,"r",encoding="utf-8",errors="ignore") as fh:
                return fh.read()
    except:
        return ""

def build_corpus_for_folders(folders, chunk_size_sentences=3):
    chunks, metas = [], []
    for folder in folders:
        if not folder or not os.path.exists(folder): continue
        files = [f for f in sorted(os.listdir(folder)) if f.lower().endswith((".pdf",".txt"))]
        for fname in files:
            p = os.path.join(folder,fname)
            text = read_file_text(p)
            if not text: continue
            sents = re.split(r'(?<=[\.\?\!])\s+',text)
            for i in range(0,max(1,len(sents)),chunk_size_sentences):
                chunk = " ".join(sents[i:i+chunk_size_sentences]).strip()
                if chunk:
                    chunks.append(chunk)
                    metas.append({"filename":fname,"folder":folder,"start":i})
    return chunks, metas

def local_search_snippets(query,chunks,metas,top_n=3):
    if not chunks: return []
    if SKLEARN_AVAILABLE:
        try:
            vectorizer = TfidfVectorizer(stop_words="english").fit(chunks+[query])
            chunk_vecs = vectorizer.transform(chunks)
            q_vec = vectorizer.transform([query])
            sims = linear_kernel(q_vec,chunk_vecs).flatten()
            top_idxs = sims.argsort()[::-1][:top_n]
            results = []
            for idx in top_idxs:
                if sims[idx]<=0: continue
                results.append({"score":float(sims[idx]),"text":chunks[idx],"meta":metas[idx]})
            return results
        except:
            pass
    out = []
    q=query.lower()
    for i,c in enumerate(chunks):
        if q in c.lower():
            out.append({"score":1.0,"text":c,"meta":metas[i]})
            if len(out)>=top_n: break
    return out

def simple_summary(text, bullets=6):
    if not text: return ""
    sents = re.split(r'(?<=[\.\?\!])\s+',text)
    selected = [s.strip() for s in sents if s.strip()][:bullets]
    return "\n".join(["- "+s for s in selected])

def model_summarize(text, bullets=6):
    if not text: return ""
    if client:
        try:
            prompt=f"Summarize into {bullets} concise bullet points:\n\n{text[:12000]}"
            resp = client.chat.completions.create(model="llama-3.3-70b-versatile",
                messages=[{"role":"user","content":prompt}],
                temperature=0.2)
            return resp.choices[0].message.content
        except:
            return simple_summary(text, bullets)
    else:
        return simple_summary(text, bullets)

def generate_audio(text):
    if not text: return ""
    if ELEVENLABS_AVAILABLE:
        try:
            elevenlabs.api_key = st.secrets.get("ELEVENLABS_API_KEY","ELEVENLABS_API_KEY_HERE")
            audio_stream = elevenlabs.generate(text=text, voice="alloy", model="eleven_multilingual_v1", stream=True)
            tmp = tempfile.NamedTemporaryFile(delete=False,suffix=".mp3")
            with open(tmp.name,"wb") as f:
                for chunk in audio_stream:
                    f.write(chunk)
            with open(tmp.name,"rb") as fh: return base64.b64encode(fh.read()).decode()
        except: pass
    if gTTS:
        try:
            tmp = tempfile.NamedTemporaryFile(delete=False,suffix=".mp3")
            gTTS(text=text, lang="en", slow=False).save(tmp.name)
            with open(tmp.name,"rb") as fh: return base64.b64encode(fh.read()).decode()
        except: pass
    return ""

# -------------------------
# Sidebar filters
# -------------------------
with st.sidebar.expander("Filters & Options", expanded=True):
    brand_options = list(brand_data.keys())
    sel_brand = st.selectbox("Brand", brand_options, index=brand_options.index(st.session_state.selected_brand))
    st.session_state.selected_brand = sel_brand
    bconf = brand_data[sel_brand]
    segment = st.selectbox("Segment", bconf["segments"])
    persona = st.selectbox("HCP Persona", bconf["personas"])
    barrier = st.multiselect("Doctor Barrier", bconf["barriers"])
    specialty = st.selectbox("Specialty", bconf["specialties"])
    objective = st.selectbox("Objective", ["Awareness","Adoption","Retention"])
    st.session_state.temperature = st.slider("Temperature",0.0,1.0,st.session_state.temperature,0.05)
    st.session_state.search_mode = st.selectbox("Search mode", ["deep","shallow"])
    st.session_state.language = st.radio("Language", ["English","Arabic"])
    if st.button("🗑️ Clear Chat"): st.session_state.chat_history=[]

with st.sidebar.expander("🌐 Add External Reference URLs (one per line)", expanded=False):
    external_urls = st.text_area("Enter URLs (one per line)").splitlines()

with st.sidebar.expander("📄 Export Options", expanded=False):
    export_format = st.radio("Choose Export Format", ["TXT","DOCX"], horizontal=True)

# -------------------------
# Title box
# -------------------------
st.markdown(f"""
<div class="title-box">
<img src="{GSK_LOGO_RAW}" class="left-logo">
<h2>💡 AI Sales Call Assistant — {brand_data[sel_brand]['display']}</h2>
<img src="{AI_LOGO_RAW}" class="right-logo">
</div>
""", unsafe_allow_html=True)

# -------------------------
# Load references and sales summaries
# -------------------------
refs_folder = bconf["references_path"]
sales_folder = bconf["sales_path"]
combined_refs = ""
if os.path.exists(refs_folder):
    for f in sorted(os.listdir(refs_folder)):
        if f.lower().endswith((".pdf",".txt")):
            combined_refs += read_file_text(os.path.join(refs_folder,f)) + "\n"
combined_sales = ""
if os.path.exists(sales_folder):
    for f in sorted(os.listdir(sales_folder)):
        if f.lower().endswith((".pdf",".txt")):
            combined_sales += read_file_text(os.path.join(sales_folder,f)) + "\n"
if not st.session_state.medical_summary and combined_refs.strip():
    st.session_state.medical_summary = model_summarize(combined_refs, bullets=6)
if not st.session_state.sales_summary and combined_sales.strip():
    st.session_state.sales_summary = model_summarize(combined_sales, bullets=6)

with st.expander("📚 Medical References Summary", expanded=False):
    st.markdown(st.session_state.medical_summary or "No medical summary available.")
with st.expander("💼 Sales Module Summary", expanded=False):
    st.markdown(st.session_state.sales_summary or "No sales summary available.")

# -------------------------
# PDF Upload and summarize
# -------------------------
uploaded_file = st.file_uploader("Upload PDF for summary", type=["pdf"])
if uploaded_file and PdfReader:
    reader = PdfReader(uploaded_file)
    pdf_text = "".join([p.extract_text() or "" for p in reader.pages])
    st.session_state.uploaded_pdf_text = pdf_text
    st.session_state.pdf_summary = model_summarize(pdf_text, bullets=6)
    st.success("PDF summarized successfully!")
if st.session_state.pdf_summary:
    with st.expander("📄 Uploaded PDF Summary", expanded=False):
        st.markdown(st.session_state.pdf_summary)

# -------------------------
# Build corpus
# -------------------------
corpus_folders = [refs_folder, sales_folder]
chunks, chunk_meta = build_corpus_for_folders(corpus_folders, chunk_size_sentences=3)

# -------------------------
# Prompt suggestions
# -------------------------
def make_suggestions(brand_key, persona_val, barriers_list, segment_val, specialty_val, objective_val):
    s=[]
    s.append(f"Generate call flow for {persona_val} focused on {objective_val}.")
    if barriers_list: s.append(f"Handle objection: {', '.join(barriers_list[:2])} for {persona_val}.")
    else: s.append(f"Identify common objections for {persona_val}.")
    s.append(f"Summarize HCP persona insights for {persona_val}.")
    s.append(f"Key talking points for {brand_data[brand_key]['display']} in {segment_val}.")
    s.append(f"Draft a short adoption message for {brand_data[brand_key]['display']} to a {specialty_val}.")
    return s

# -------------------------
# AI Response with APACT + humanized + interactive feedback
# -------------------------
def add_ai_response(prompt, follow_up=False):
    snippets = local_search_snippets(prompt, chunks, chunk_meta, top_n=5)
    citation = "\n".join([f"{s['meta']['filename']} ({s['score']:.2f})" for s in snippets])

    response_lines = []

    if not follow_up:
        # --- APACT ---
        response_lines.append(f"**Acknowledge:** Thank you for raising this concern. I understand your perspective.")
        response_lines.append("**Probing:** Could you clarify if your main concern is about efficacy, safety, or patient eligibility?")
        response_lines.append("**Actions:** Based on your input, here are recommended steps:")
        for step in bconf["call_flow"]:
            step_snippets = [s['text'] for s in snippets if step.lower() in s['text'].lower()]
            if step_snippets:
                response_lines.append(f"**{step}:**")
                for sn in step_snippets:
                    response_lines.append(f"- {sn}")
            else:
                response_lines.append(f"**{step}:** - Refer to the sales module and uploaded references for guidance.")
        response_lines.append("**Confirm:** Does this approach address your concern sufficiently?")
        response_lines.append("**Transition:** If yes, we can move on to the next discussion point or objective.")
        response_lines.append("\n*Note: Tailored using sales module and uploaded references.*")
    else:
        # Follow-up for feedback
        response_lines.append("I noticed you disliked the previous answer. Could you help me understand better?")
        response_lines.append("- What specific part was unclear or insufficient?")
        response_lines.append("- Are you looking for more examples, data, or step-by-step guidance?")
        response_lines.append("- Any particular objection you want me to focus on next?")

    ai_text = "\n".join(response_lines)
    st.session_state.chat_history.append({"role":"assistant","content":ai_text,"citation":citation})

# -------------------------
# Chat container and input
# -------------------------
chat_container = st.container()

with st.expander("💡 Prompt Suggestions (Click to Expand)", expanded=False):
    suggs = make_suggestions(sel_brand, persona, barrier, segment, specialty, objective)
    sugg_cols = st.columns(3)
    for i, s in enumerate(suggs):
        col = sugg_cols[i % 3]
        if col.button(s, key=f"sugg_{i}"):
            st.session_state.main_input = s

with st.form("main_input_form", clear_on_submit=True):
    user_input = st.text_area("Ask something:", st.session_state.main_input, height=72)
    submitted = st.form_submit_button("Send")
    if submitted and user_input.strip():
        st.session_state.chat_history.append({"role":"user","content":user_input.strip()})
        add_ai_response(user_input.strip())
        st.session_state.main_input = ""

# -------------------------
# Display chat with audio and interactive feedback
# -------------------------
with chat_container:
    for idx,entry in enumerate(st.session_state.chat_history):
        if entry["role"]=="user":
            st.markdown(f'<div class="chat-bubble-user">{escape(entry["content"])}</div>',unsafe_allow_html=True)
        else:
            st.markdown(f'<div class="chat-bubble-ai">{escape(entry["content"])}</div>',unsafe_allow_html=True)
            if "citation" in entry and entry["citation"]:
                st.markdown(f'<div class="citation-box">{escape(entry["citation"])}</div>',unsafe_allow_html=True)
            # Audio playback
            audio_b64 = generate_audio(entry["content"])
            if audio_b64:
                st.audio(io.BytesIO(base64.b64decode(audio_b64)), format="audio/mp3")
            # Feedback buttons
            fb_cols = st.columns(3)
            if entry["content"] not in st.session_state.feedback:
                if fb_cols[0].button("👍 Like", key=f"like_{idx}"): st.session_state.feedback[entry["content"]]="like"
                if fb_cols[1].button("👎 Dislike", key=f"dislike_{idx}"): 
                    st.session_state.feedback[entry["content"]]="dislike"
                    add_ai_response("Follow-up based on user dislike", follow_up=True)
                if fb_cols[2].button("ℹ️ Need More", key=f"needmore_{idx}"): 
                    st.session_state.feedback[entry["content"]]="need_more"
                    add_ai_response("The user requested more information; expand the previous answer.", follow_up=True)

# -------------------------
# Footer disclaimer
# -------------------------
st.markdown("""
<div class="fixed-disclaimer">
💡 This tool is for internal sales support purposes only. All medical info should be verified from official sources. 
</div>
""",unsafe_allow_html=True)
