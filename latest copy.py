# app.py
import os
import re
import time
import base64
import tempfile
import asyncio
from typing import Optional
from datetime import datetime
from io import BytesIO

import streamlit as st
from PIL import Image, ImageStat
import requests
import PyPDF2
import edge_tts
import html 
# Groq client (optional)
try:
    import groq
    from groq import Groq
except Exception:
    Groq = None

# Optional docx export
try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# ----------------------------
# Page config
# ----------------------------
st.set_page_config(page_title="AI Sales Call Assistant", page_icon="💡", layout="wide")

# ----------------------------
# GROQ API Key (preferred via environment variable)
# ----------------------------
# Set your GROQ_API_KEY in the environment for safety:
# export GROQ_API_KEY="your_real_key"
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "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")
client = Groq(api_key=GROQ_API_KEY) if (GROQ_API_KEY and Groq is not None and "PUT_YOUR" not in GROQ_API_KEY) else None
if client is None:
    st.info("Groq AI client not configured. Set environment variable GROQ_API_KEY with a valid key to enable AI summarization/answers.")

# ----------------------------
# Session defaults
# ----------------------------
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "uploaded_pdf_text" not in st.session_state:
    st.session_state.uploaded_pdf_text = ""
if "extracted_medical_ref" not in st.session_state:
    st.session_state.extracted_medical_ref = ""
if "pdf_summary" not in st.session_state:
    st.session_state.pdf_summary = ""
if "language" not in st.session_state:
    st.session_state.language = "English"
if "voice_pref" not in st.session_state:
    st.session_state.voice_pref = "English Neural"

# ----------------------------
# Assets & variables
# ----------------------------
BACKGROUND_URL = "https://www.shutterstock.com/image-photo/excited-girl-white-shirt-using-260nw-708132598.jpg"
GSK_LOGO_URL = "https://i-cf65.gskstatic.com/content/dam/cf-pharma/gskusmedicalaffairs/en_US/logos/gsk-logo-white.png?auto=format"

def get_brightness(url: str) -> int:
    try:
        r = requests.get(url, timeout=6)
        img = Image.open(BytesIO(r.content)).convert("L")
        stat = ImageStat.Stat(img)
        return int(stat.mean[0])
    except Exception:
        return 255

brightness = get_brightness(BACKGROUND_URL)
text_color = "black" if brightness > 130 else "white"

# ----------------------------
# Styling (use .format to safely inject URL)
# ----------------------------
CSS = """
<style>
/* App background container (observed by JS to resize) */
[data-testid="stAppViewContainer"] {{
  background-image: url("{bg}");
  background-repeat: no-repeat;
  background-position: right top;
  background-attachment: fixed;
  background-size: auto 150%;
  transition: background-size 0.25s ease;
}}

/* Sidebar style and visible control borders */
[data-testid="stSidebar"] > div:first-child {{
  background: #ddddd;
  padding: 12px;
}}
[data-testid="stSidebar"] .stSelectbox, [data-testid="stSidebar"] .stMultiselect,
[data-testid="stSidebar"] .stRadio, [data-testid="stSidebar"] .stCheckbox,
[data-testid="stSidebar"] .stFileUploader {{
  border: 1px solid #e6e6e6;
  border-radius: 10px;
  padding: 8px;
  margin-bottom: 10px;
  background-color: #fff;
}}

/* Title box */
.title-box {{
  background: rgba(245,245,245,0.6);
  padding: 22px;
  border-radius: 14px;
  text-align: center;
  max-width: 80%;
  margin: 12px auto;
}}
.title-box h1 {{ margin:0; font-size:34px; font-weight:800; color:#000; }}
.title-box p {{ margin:6px 0 0 0; font-size:16px; color:#333; }}

/* PDF summary box */
.pdf-summary-box {{
  background: rgba(245,245,245,0.8);
  padding: 12px;
  border-radius: 14px;
  border: 1px solid #000;
  margin-bottom: 12px;
}}

/* Chat container & bubbles */
.chat-container {{
  height: 56vh;
  overflow:auto;
  padding:12px;
  border-radius:10px;
  background: rgba(255,255,255,0.76);
}}
.chat-bubble-user, .chat-bubble-ai {{
  display:inline-block;
  padding:12px;
  border-radius:12px;
  margin:8px 0;
  max-width: 86%;
  word-wrap: break-word;
  color: #000;
}}
.chat-bubble-user {{ background: #eef9e6; margin-left:auto; }}
.chat-bubble-ai {{ background: #f5f7fa; margin-right:auto; }}

/* Inline PDF snippet inside AI bubble */
.pdf-summary-inline {{
  margin-top:8px;
  background: rgba(245,245,245,0.7);
  padding:10px;
  border-radius:8px;
  border:1px solid #1111;
}}

/* Bottom fixed input */
.bottom-bar {{
  position: fixed;
  bottom: 12px;
  left: 16px;
  right: 16px;
  z-index: 1200;
  background: rgba(255,255,255,0.98);
  padding:10px;
  border-radius:12px;
  box-shadow: 0 6px 18px rgba(0,0,0,0.06);
  display:flex;
  gap:12px;
  align-items:center;
}}
.bottom-bar input[type="text"] {{
  flex:1;
  padding:10px 12px;
  border-radius:8px;
  border:1px solid #ddd;
  outline:none;
}}
.bottom-bar button {{
  min-width:110px;
  padding:8px 12px;
  border-radius:8px;
  background:#ff8c00;
  color:white;
  border:none;
  font-weight:600;
  cursor:pointer;
}}

/* small screen tweaks */
@media (max-width: 430px) {{
  .title-box h1 {{ font-size:22px; }}
  .chat-container {{ height:48vh; }}
  .bottom-bar {{ left:8px; right:8px; bottom:8px; padding:8px; }}
}}

/* highlight classes used inside AI bubble */
.highlight-step {{ font-weight:700; color:#000; }}
.highlight-figure {{ font-weight:700; color:#d35400; }}
</style>
""".format(bg=BACKGROUND_URL)

st.markdown(CSS, unsafe_allow_html=True)

# ----------------------------
# JS: observe sidebar expand/collapse and resize background
# ----------------------------
SIDEBAR_JS = """
<script>
(function(){
  const sidebar = document.querySelector('[data-testid="stSidebar"]');
  const app = document.querySelector('[data-testid="stAppViewContainer"]');
  if(!sidebar || !app) return;
  function updateBg(){
    const expanded = sidebar.getAttribute('aria-expanded') === 'true';
    app.style.backgroundSize = expanded ? 'auto 80%' : 'auto 100%';
  }
  updateBg();
  const mo = new MutationObserver(updateBg);
  mo.observe(sidebar, { attributes: true });
})();
</script>
"""
st.markdown(SIDEBAR_JS, unsafe_allow_html=True)

# small JS to autoscroll chat container
SCROLL_JS = """
<script>
function scrollChat(){
  const el = document.querySelector('.chat-container');
  if (el) el.scrollTop = el.scrollHeight;
}
setTimeout(scrollChat, 200);
</script>
"""

# top logo + title
st.markdown(f'<div style="position:auto; top:80px; left:18px; z-index:1200;"><img src="{GSK_LOGO_URL}" width="140" /></div>', unsafe_allow_html=True)
st.markdown('<div class="title-box"><h1>💡 AI Sales Call Assistant</h1><p>Powered by AI to equip sales reps for smarter HCP conversations</p></div>', unsafe_allow_html=True)
st.markdown('<p style="text-align:center;font-weight:600;">⚠️ Disclaimer: For training and educational purposes only.</p>', unsafe_allow_html=True)

# ----------------------------
# Data lists
# ----------------------------
gsk_brands = {
    "Shingrix": "https://www.shingrix.com/",
    "Trelegy": "https://www.trelegy.com/",
    "Zejula": "https://www.zejula.com/"
}
gsk_brands_images = {
    "Shingrix":"https://www.oma-apteekki.fi/WebRoot/NA/Shops/na/67D6/48DA/D0B0/D959/ECAF/0A3C/0E02/D573/3ad67c4e-e1fb-4476-a8a0-873423d8db42_3Dimage.png",
    "Trelegy":"https://www.1uphealth.com/wp-content/uploads/2020/11/trelegy.png",
    "Zejula":"https://cdn.salla.sa/QeZox/eyy7B0bg8D7a0Wwcov6UshWFc04R6H8qIgbfFq8u.png"
}

race_segments = ["R – Reach", "A – Acquisition", "C – Conversion", "E – Engagement"]
doctor_barriers = [
    "HCP does not consider HZ a risk","No time for discussion","Cost concerns",
    "Not convinced of efficacy","Accessibility/Logistics","Patient reluctance","Other clinical doubts"
]
personas = ["Uncommitted Vaccinator","Reluctant Efficiency","Patient Influenced","Committed Vaccinator"]
sales_call_flow = [
    "Prepare the call",
    "Engage",
    "Create opportunities",
    "Impact GSO (Good sell outcome)",
    "Influence",
    "Analyze and post call analysis"
]
APACT_STEPS = ["Acknowledge","Probing","Action","Confirm","Transition"]
objectives = ["Awareness","Adoption","Retention"]
specialties = ["GP","Cardiologist","Dermatologist","Endocrinologist","Pulmonologist","Rheumatologist","Internal Medicine","Diabetologist","Neurologist","Pneumologist"]

# ----------------------------
# Helper functions (defined early to avoid NameError)
# ----------------------------
def highlight_content_for_display(content: str) -> str:
    """Wrap call steps, APACT steps, and percentages in span classes (used for AI HTML rendering)."""
    if not content:
        return content
    # Bold/Highlight call steps
    for step in sales_call_flow:
        content = re.sub(rf"\b{re.escape(step)}\b", f"<span class='highlight-step'>{step}</span>", content)
    # APACT
    for ap in APACT_STEPS:
        content = re.sub(rf"\b{re.escape(ap)}\b", f"<span class='highlight-step'>{ap}</span>", content)
    # Percent figures
    for fig in re.findall(r"\d+\.?\d*%", content):
        content = content.replace(fig, f"<span class='highlight-figure'>{fig}</span>")
    return content

def build_ai_bubble_content(ai_text: str, inject_pdf_lines: int = 6) -> str:
    text = (ai_text or "").replace("\n", "<br>")
    text = highlight_content_for_display(text)
    pdf_html = ""
    if st.session_state.pdf_summary:
        # take a few bullet lines from PDF summary to show inline
        pdf_lines = [ln.strip() for ln in st.session_state.pdf_summary.splitlines() if ln.strip()]
        sample = pdf_lines[:inject_pdf_lines]
        if sample:
            pdf_html = "<div class='pdf-summary-inline'>" + "<br>".join([f"- {ln}" for ln in sample]) + ( "<br>..." if len(pdf_lines) > inject_pdf_lines else "" ) + "</div>"
    return f"{text}{pdf_html}"

def render_chat_history():
    html_out = ""
    for msg in st.session_state.chat_history:
        content = html.escape(msg["content"])  # safe escaping
        if msg["role"] == "user":
            html_out += f"<div class='chat-bubble-user'>{content}</div>"
        else:
            ai_bubble = build_ai_bubble_content(msg["content"])
            audio_html = ""
            if msg.get("audio"):
                audio_html = f"<br><audio controls style='margin-top:8px;'><source src='data:audio/mp3;base64,{msg['audio']}' type='audio/mp3'></audio>"
            html_out += f"<div class='chat-bubble-ai'>{ai_bubble}{audio_html}</div>"
    st.markdown(html_out, unsafe_allow_html=True)

async def _edge_save_async(ssml_text: str, voice: str, outpath: str):
    comm = edge_tts.Communicate(ssml_text, voice=voice)
    await comm.save(outpath)

def synthesize_tts_humanized(text: str, lang: str, voice_pref: str) -> Optional[str]:
    """
    Create humanized speech:
      - converts '45%' -> '45 percent'
      - strips problematic symbols
      - inserts small breaks (SSML)
    Returns base64-encoded mp3 or None on failure.
    """
    if not text:
        return None
    # Replace percent signs to speech-friendly words
    text = re.sub(r'(\d+)\s*%', r'\1 percent', text)
    # Remove characters that make TTS choppy (but keep sentence punctuation .?!)
    text = re.sub(r'[;:{}\[\]\*\^<>@#\$%&\|~_/\\+]', '', text)
    text = re.sub(r'\s+', ' ', text).strip()
    # Split into sentences preserving ends for pauses
    sentences = re.split(r'(?<=[.?!])\s+', text)
    ssml_parts = []
    for s in sentences:
        s = s.strip()
        if not s:
            continue
        # Slightly longer break after end of sentence
        ssml_parts.append(f"<prosody rate='medium'>{s}<break time='0.45s'/></prosody>")
    ssml = "<speak>" + " ".join(ssml_parts) + "</speak>"

    # choose voice
    if voice_pref == "Arabic Neural":
        voice = "ar-EG-SalmaNeural"
    elif voice_pref == "English Neural":
        voice = "en-US-AriaNeural"
    else:
        voice = "en-US-AriaNeural"

    tmp = tempfile.NamedTemporaryFile(suffix=".mp3", delete=False)
    tmp_name = tmp.name
    tmp.close()
    try:
        asyncio.run(_edge_save_async(ssml, voice, tmp_name))
        with open(tmp_name, "rb") as f:
            data = f.read()
        return base64.b64encode(data).decode("utf-8")
    except Exception as e:
        st.warning(f"TTS generation failed: {e}")
        return None
    finally:
        try:
            if os.path.exists(tmp_name):
                os.remove(tmp_name)
        except Exception:
            pass

def build_prompt(user_input: str, language: str) -> str:
    """Build the assistant prompt, prioritizing PDF summary & references and enforcing call flow & APACT."""
    pdf_summary = st.session_state.pdf_summary or ""
    refs = st.session_state.extracted_medical_ref or "None"
    instructions = [
        "- Use the uploaded PDF summary and extracted medical references as primary sources for clinical claims when applicable.",
        "- Cite references when available (e.g., CDC, NEJM).",
        "- Provide a short 'Sample script' (3–6 lines) for the rep to say.",
        "- Match requested tone and length."
    ]
    # if user asks for call flow specifically
    if re.search(r"\b(sales call flow|call flow|sales flow|sales steps)\b", user_input, flags=re.I):
        steps_text = ", ".join([f"**{s}**" for s in sales_call_flow])
        instructions.append(f"When asked for 'sales call flow', return the steps as bold bullet points in this exact order: {steps_text}. Provide 1-2 short action lines per step.")
    # objections -> use APACT
    if re.search(r"\b(objection|concern|barrier|hesitat|not convinced|resist)\b", user_input, flags=re.I):
        instructions.append("When handling objections, structure the response using APACT: **Acknowledge**, **Probing**, **Action**, **Confirm**, **Transition**. Bold the APACT step titles.")
    # always ask to bold steps and figures
    instructions.append("Bold sales steps, APACT titles, and any notable numeric figures (e.g., 45%).")

    prompt_parts = [
        f"Language: {language}",
        f"User input: {user_input}",
        f"Brand: {brand}",
        f"RACE Segment: {segment}",
        f"Doctor Barrier(s): {', '.join(barrier) if barrier else 'None'}",
        f"Objective: {objective}",
        f"Doctor Specialty: {specialty}",
        f"HCP Persona: {persona}",
        "",
        "Instructions:",
        *instructions,
        "",
        "PDF Summary (if any):",
        pdf_summary or "None",
        "",
        "Extracted references:",
        refs,
        "",
        f"Response tone: {response_tone}. Desired length: {response_length}."
    ]
    return "\n".join(prompt_parts)

def call_groq_with_retry(prompt: str, language: str, max_retries: int = 3, base_delay: int = 2) -> str:
    """Call Groq chat completion with simple retry/fallback logic."""
    if client is None:
        return "⚠️ AI service not configured. Set GROQ_API_KEY in environment to enable AI generation."
    models = [
        "meta-llama/llama-4-scout-17b-16e-instruct",
        "meta-llama/llama-4-scout-13b-instruct",
    ]
    last_err = None
    for model in models:
        for attempt in range(1, max_retries + 1):
            try:
                resp = client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": f"You are a helpful sales assistant that responds in {language}."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.65,
                )
                return resp.choices[0].message.content
            except Exception as e:
                last_err = e
                err_msg = str(e).lower()
                if "over capacity" in err_msg or "503" in err_msg or "internal_server_error" in err_msg:
                    wait = base_delay * (2 ** (attempt - 1))
                    st.warning(f"Model {model} busy. Retrying in {wait}s (attempt {attempt}/{max_retries})...")
                    time.sleep(wait)
                    continue
                if "authentication" in err_msg or "unauthorized" in err_msg:
                    return "⚠️ Authentication error with Groq. Please check your GROQ_API_KEY."
                return f"⚠️ Error generating response: {e}"
        st.info(f"Switching to fallback model after {max_retries} attempts for {model}.")
    return f"⚠️ AI call failed after retries. Last error: {last_err}"

# ----------------------------
# Sidebar (filters) - runs early so variables exist for prompt builder
# ----------------------------
with st.sidebar.expander("Filters & Options", expanded=True):
    st.markdown('<div style="font-weight:800; margin-bottom:8px;">Filters & Options</div>', unsafe_allow_html=True)
    brand = st.selectbox("Select Brand / اختر العلامة التجارية", options=list(gsk_brands.keys()))
    img_path = gsk_brands_images.get(brand)
    if img_path:
        try:
            resp = requests.get(img_path, timeout=6)
            img = Image.open(BytesIO(resp.content))
            st.image(img, width=160)
        except Exception:
            st.image("https://via.placeholder.com/160x90.png?text=No+Image", width=160)

    segment = st.selectbox("Select RACE Segment / اختر شريحة RACE", options=race_segments)
    barrier = st.multiselect("Select Doctor Barrier / اختر حاجز الطبيب", options=doctor_barriers, default=[])
    objective = st.selectbox("Select Objective / اختر الهدف", options=objectives)
    specialty = st.selectbox("Select Doctor Specialty / اختر تخصص الطبيب", options=specialties)
    persona = st.selectbox("Select HCP Persona / اختر شخصية الطبيب", options=personas)
    response_length = st.selectbox("Response Length / اختر طول الرد", ["Short", "Medium", "Long"])
    response_tone = st.selectbox("Response Tone / اختر نبرة الرد", ["Formal", "Casual", "Friendly", "Persuasive"])
    st.session_state.language = st.radio("Language / اختر اللغة", options=["English", "العربية"], index=0, horizontal=True)
    st.session_state.voice_pref = st.selectbox("Voice preference", ["English Neural", "Arabic Neural", "Default"])

# ----------------------------
# PDF Upload & Summarization
# ----------------------------
st.markdown("### 📄 Upload Medical Reference PDF")
uploaded_pdf = st.file_uploader("Upload PDF for AI reference", type=["pdf"])
if uploaded_pdf:
    try:
        reader = PyPDF2.PdfReader(uploaded_pdf)
        full_text = "".join([p.extract_text() or "" for p in reader.pages])
        st.session_state.uploaded_pdf_text = full_text[:2000] + "..." if len(full_text) > 2000 else full_text
        # extract likely references heuristically
        matches = re.findall(r"(?:CDC|FDA|Guideline|Study|Journal|20\d{2}|Lancet|NEJM|BMJ|JAMA)[^.\n]*", full_text, flags=re.I)
        st.session_state.extracted_medical_ref = ", ".join(matches) if matches else "None"
        st.success("✅ PDF processed (extracted text and references).")

        # chunk & summarize using Groq if configured
        if client is None:
            st.warning("Groq client not configured: PDF auto-summarize unavailable. Set GROQ_API_KEY in environment.")
            st.session_state.pdf_summary = ""
        else:
            chunk_size = 5000
            chunks = [full_text[i:i+chunk_size] for i in range(0, len(full_text), chunk_size)]
            summaries = []
            for chunk in chunks:
                summary_prompt = (
                    "You are a concise medical summarizer for sales reps. Produce short bullet points with key results, "
                    "practical recommendations, and notable figures (write % as 'percent'). Keep it actionable and short.\n\n"
                ) + chunk[:6000]
                try:
                    resp = client.chat.completions.create(
                        model="meta-llama/llama-4-scout-17b-16e-instruct",
                        messages=[{"role":"system","content":"You are a concise medical summarizer."},{"role":"user","content":summary_prompt}],
                        temperature=0.2,
                    )
                    summaries.append(resp.choices[0].message.content)
                except Exception as e:
                    st.warning(f"Chunk summarization warning: {e}")
                    # continue with remaining chunks
            st.session_state.pdf_summary = "\n".join(summaries).strip()

        # Show PDF summary inside an expander
        if st.session_state.pdf_summary:
            with st.expander("📑 PDF Summary (expand/collapse)", expanded=False):
                # render as bullet list
                lines = [ln.strip() for ln in st.session_state.pdf_summary.splitlines() if ln.strip()]
                md = "\n".join([f"- {ln}" for ln in lines])
                st.markdown(f'<div class="pdf-summary-box">{st.markdown(md, unsafe_allow_html=False) or ""}</div>', unsafe_allow_html=True)
        if st.session_state.extracted_medical_ref:
            st.info(f"📚 Extracted references: {st.session_state.extracted_medical_ref}")
    except Exception as e:
        st.error("PDF error: " + str(e))

# ----------------------------
# Chat area rendering
# ----------------------------
st.markdown("<h3>💬 Chat</h3>", unsafe_allow_html=True)
render_chat_history()

# ----------------------------
# Bottom input form (fixed)
# ----------------------------
with st.form("bottom_chat_form", clear_on_submit=True):
    user_text = st.text_input("Type your message... / اكتب رسالتك هنا", key="bottom_prompt_input", label_visibility="collapsed")
    submitted = st.form_submit_button("Send")

if submitted and user_text and user_text.strip():
    # append user message
    st.session_state.chat_history.append({"role": "user", "content": user_text.strip(), "time": datetime.now().strftime("%H:%M")})
    render_chat_history()

    # build prompt and call AI
    prompt = build_prompt(user_text.strip(), st.session_state.language)
    ai_output = call_groq_with_retry(prompt, st.session_state.language)
    # generate TTS (best-effort)
    audio_b64 = synthesize_tts_humanized(ai_output, st.session_state.language, st.session_state.voice_pref)
    # append AI message
    st.session_state.chat_history.append({"role": "ai", "content": ai_output, "time": datetime.now().strftime("%H:%M"), "audio": audio_b64})
    render_chat_history()

# ----------------------------
# Quick convenience: show sales call flow button
# ----------------------------
if st.button("Show GSK Sales Call Flow"):
    # Build a quick preformatted AI-style response (bold steps)
    flow_lines = [f"**{s}**: Provide 1–2 lines of practical guidance for this step." for s in sales_call_flow]
    flow_text = "\n\n".join(flow_lines)
    st.session_state.chat_history.append({"role": "ai", "content": flow_text, "time": datetime.now().strftime("%H:%M"), "audio": None})
    render_chat_history()

# ----------------------------
# Bottom controls: clear + export
# ----------------------------
cols = st.columns([1, 1])
with cols[0]:
    if st.button("🗑️ Clear Chat / مسح المحادثة"):
        st.session_state.chat_history = []
        st.session_state.uploaded_pdf_text = ""
        st.session_state.extracted_medical_ref = ""
        st.session_state.pdf_summary = ""
        st.experimental_rerun()
with cols[1]:
    if DOCX_AVAILABLE and st.session_state.chat_history:
        if st.button("📥 Export Chat (.docx)"):
            doc = Document()
            doc.add_heading("AI Sales Call Assistant Chat History", 0)
            for msg in st.session_state.chat_history:
                role = "User" if msg.get("role") == "user" else "AI"
                doc.add_paragraph(f"{role} [{msg.get('time','')}]")
                text_content = re.sub(r'<.*?>', '', msg.get("content", ""))
                doc.add_paragraph(text_content)
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            doc.save(tmp.name)
            tmp.close()
            with open(tmp.name, "rb") as f:
                data = f.read()
            st.download_button("⬇️ Download Chat History (.docx)", data=data, file_name="chat_history.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# footer: brand leaflet link
st.markdown(f"[📄 Brand Leaflet - {brand}]({gsk_brands.get(brand)})")





# app.py
import streamlit as st
from PIL import Image
from io import BytesIO
import re
import tempfile
import base64
from groq import Groq
from PyPDF2 import PdfReader
from html import escape

# Optional docx export
try:
    from docx import Document
    DOCX_AVAILABLE = True
except:
    DOCX_AVAILABLE = False

# ---------------------------- TTS Setup ----------------------------
try:
    import elevenlabs
    ELEVENLABS_AVAILABLE = True
except ModuleNotFoundError:
    ELEVENLABS_AVAILABLE = False

from gtts import gTTS

# ElevenLabs config
ELEVENLABS_API_KEY = st.secrets.get("ELEVENLABS_API_KEY", "")
ELEVENLABS_VOICE_ID = st.secrets.get("ELEVENLABS_VOICE_ID", "")
if ELEVENLABS_AVAILABLE and ELEVENLABS_API_KEY and ELEVENLABS_VOICE_ID:
    elevenlabs.api_key = ELEVENLABS_API_KEY
else:
    ELEVENLABS_AVAILABLE = False

def generate_audio(text):
    for step in ["Acknowledge","Probing","Action","Confirm","Transition"]:
        text = text.replace(step, f"{step} ...")
    text = re.sub(r'[.,*]', '', text)
    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3")
    if ELEVENLABS_AVAILABLE:
        audio_stream = elevenlabs.generate(text=text, voice=ELEVENLABS_VOICE_ID, stream=True)
        with open(tmp_file.name, "wb") as f:
            for chunk in audio_stream:
                f.write(chunk)
    else:
        tts = gTTS(text=text, lang="en", slow=True)
        tts.save(tmp_file.name)
    with open(tmp_file.name, "rb") as f:
        audio_bytes = f.read()
        audio_base64 = base64.b64encode(audio_bytes).decode()
    return audio_base64

# ---------------------------- CONFIG ----------------------------
st.set_page_config(page_title="GSK AI Sales Call Assistant", layout="wide")

# ---------------------------- Session Defaults ----------------------------
if "chat_history" not in st.session_state or not isinstance(st.session_state.chat_history, list):
    st.session_state.chat_history = []
if "uploaded_pdf_text" not in st.session_state:
    st.session_state.uploaded_pdf_text = ""
if "pdf_summary" not in st.session_state:
    st.session_state.pdf_summary = ""
if "voice_pref" not in st.session_state:
    st.session_state.voice_pref = "Old Male"
if "language" not in st.session_state:
    st.session_state.language = "English"
if "pdf_search_keyword" not in st.session_state:
    st.session_state.pdf_search_keyword = ""
if "pdf_summary_size" not in st.session_state:
    st.session_state.pdf_summary_size = "Normal"

# ---------------------------- Assets ----------------------------
BACKGROUND_URL = "https://sdmntprukwest.oaiusercontent.com/files/00000000-abd4-6243-82cf-168367664603/raw?se=2025-10-02T08%3A55%3A08Z&sp=r&sv=2024-08-04&sr=b&scid=da9b1fe8-d683-5331-8dac-5d17ac775ed0&skoid=82a3371f-2f6c-4f81-8a78-2701b362559b&sktid=a48cca56-e6da-484e-a814-9c849652bcb3&skt=2025-10-02T05%3A07%3A48Z&ske=2025-10-03T05%3A07%3A48Z&sks=b&skv=2024-08-04&sig=zev17ijVwaJyIwxogpGkQRRHoIWzd7z4Ic%2BWeVhPdjc%3D"
GSK_LOGO_URL = "https://usppg.org/wp-content/uploads/2025/04/GSK-logo.png"
AI_LOGO_URL = "https://sdmntpritalynorth.oaiusercontent.com/files/00000000-42e0-6246-8bd4-812f66b46668/raw?se=2025-10-02T09%3A09%3A04Z&sp=r&sv=2024-08-04&sr=b&scid=04001bb8-a622-5394-8e9b-f0e7f4f6f1f2&skoid=82a3371f-2f6c-4f81-8a78-2701b362559b&sktid=a48cca56-e6da-484e-a814-9c849652bcb3&skt=2025-10-02T04%3A32%3A34Z&ske=2025-10-03T04%3A32%3A34Z&sks=b&skv=2024-08-04&sig=eStxlnunHXrvS6s65lQTrZCH1ziJhQ6mUxgpbnT/zeY%3D"

# ---------------------------- CSS ----------------------------
CSS = f"""
<style>
/* Background */
[data-testid="stAppViewContainer"] {{
  background-image: url("{BACKGROUND_URL}");
  background-repeat: no-repeat;
  background-position: right top;
  background-attachment: fixed;
  background-size: auto 150%;
}}

/* Title box */
.title-box {{
  background: rgba(245,245,245,0.85);
  padding: 15px;
  border-radius: 16px;
  text-align: center;
  margin: 12px auto;
  width: 650px;
  position: relative;
}}
.title-box img.ai-logo {{
    position: absolute;
    top: 10px;
    right: 10px;
    width: 80px;
}}

/* PDF summary box */
.pdf-summary-box {{
  background: #E6F0FF; 
  padding: 12px; 
  border-radius: 14px; 
  margin-bottom: 12px;
  white-space: pre-line;
}}

/* Chat area */
.chat-container {{
  max-height: 65vh;
  overflow-y: auto;
  padding: 12px;
  border-radius: 10px;
  background: rgba(255,255,255,0.85);
  margin-bottom: 20px;
}}

/* Bubbles */
.chat-bubble-user, .chat-bubble-ai, .chat-bubble-audio {{
  display:block;
  padding:12px;
  border-radius:12px;
  margin:12px 0;
  max-width: 86%;
  word-wrap: break-word;
}}
.chat-bubble-user {{ background: #0078D7; color:white; margin-left:auto; }}
.chat-bubble-ai {{ background: #d9f0ff; margin-right:auto; color:#000; }}
.chat-bubble-audio {{ background: #e2e2e2; margin-right:auto; font-size:0.9em; padding:10px; margin-top:12px; }}

/* Fixed chat input at bottom */
.fixed-chat-input {{
    position: fixed;
    bottom: 20px;
    left: 20px;
    right: 20px;
    z-index: 10002;
}}
.fixed-chat-input textarea {{
    width: 100%;
    min-height: 60px;
    max-height: 180px;
    resize: vertical;
}}
.send-button {{
    position: fixed;
    bottom: 20px;
    right: 30px;
    z-index: 10003;
    height: 40px;
    width: 100px;
}}
</style>
"""
st.markdown(CSS, unsafe_allow_html=True)

# ---------------------------- GROQ Client ----------------------------
GROQ_API_KEY = "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"
client = Groq(api_key=GROQ_API_KEY)

# ---------------------------- Filters / Sidebar ----------------------------
gsk_brands = ["Shingrix", "Trelegy", "Zejula"]
race_segments = ["R – Reach", "A – Acquisition", "C – Conversion", "E – Engagement"]
doctor_barriers = ["HCP does not consider HZ a risk","No time for discussion","Cost concerns","Not convinced of efficacy"]
personas = ["Uncommitted Vaccinator","Reluctant Efficiency","Patient Influenced","Committed Vaccinator"]
sales_call_flow = ["Prepare","Engage","Create Opportunities","Impact GSO","Influence","Post Call Analysis"]
APACT_STEPS = ["Acknowledge","Probing","Action","Confirm","Transition"]
objectives = ["Awareness","Adoption","Retention"]
specialties = ["GP","Cardiologist","Dermatologist","Endocrinologist"]

with st.sidebar.expander("Filters & Options", expanded=True):
    brand = st.selectbox("Brand", gsk_brands, key="select_brand")
    specialty = st.selectbox("Specialty", specialties, key="select_specialty")
    segment = st.selectbox("RACE Segment", race_segments, key="select_segment")
    persona = st.selectbox("HCP Persona", personas, key="select_persona")
    barrier = st.multiselect("Doctor Barrier", doctor_barriers, key="select_barrier")
    objective = st.selectbox("Objective", objectives, key="select_objective")
    response_tone = st.selectbox("Response Tone", ["Formal","Casual","Friendly","Persuasive"], key="select_tone")
    response_length = st.selectbox("Response Length", ["Short","Medium","Long"], key="select_length")
    st.session_state.language = st.radio("Language", ["English","Arabic"], horizontal=True, key="select_language")

# ---------------------------- Title Box ----------------------------
st.markdown(f'''
<div class="title-box">
    <img src="{GSK_LOGO_URL}" width="140">
    <img src="{AI_LOGO_URL}" class="ai-logo">
    <h1>💡 AI Sales Call Assistant</h1>
    <p>Powered by AI to equip reps for smarter HCP conversations</p>
</div>
''', unsafe_allow_html=True)

# ---------------------------- PDF Upload & Summary ----------------------------
with st.expander("📄 PDF Summary", expanded=False):
    uploaded_pdf = st.file_uploader("Upload PDF for AI reference", type=["pdf"])
    st.session_state.pdf_summary_size = st.radio("PDF Summary Size", ["Consisted","Normal","Detailed"], horizontal=True)
    if uploaded_pdf:
        reader = PdfReader(uploaded_pdf)
        full_text = "".join([p.extract_text() or "" for p in reader.pages])
        st.session_state.uploaded_pdf_text = full_text
        bullets_count = {"Consisted":5,"Normal":10,"Detailed":20}.get(st.session_state.pdf_summary_size,10)
        try:
            summary_prompt = f"Summarize the document into {bullets_count} bullet points:\n{full_text[:12000]}"
            ai_summary = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role":"system","content":"You are a helpful assistant."},
                          {"role":"user","content":summary_prompt}],
                temperature=0.4
            )
            st.session_state.pdf_summary = ai_summary.choices[0].message.content
        except Exception:
            fallback_bullets = re.findall(r'([A-Z][^.]{20,200})', full_text)
            st.session_state.pdf_summary = "\n".join(fallback_bullets[:bullets_count])
    if st.session_state.pdf_summary:
        st.markdown(f'<div class="pdf-summary-box">{escape(st.session_state.pdf_summary)}</div>', unsafe_allow_html=True)

# ---------------------------- AI Chat ----------------------------
def generate_ai_response(user_input):
    if "sales call flow" in user_input.lower():
        prompt = f"Build a sales call flow for {persona} using RACE steps."
    elif "handle objection" in user_input.lower() or "apact" in user_input.lower():
        prompt = f"Generate APACT approach steps for the objection: {user_input}"
    else:
        prompt = f"Answer the query professionally: {user_input}"

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role":"system","content":"You are a helpful sales assistant AI."},
                  {"role":"user","content":prompt}],
        temperature=0.65
    )
    return response.choices[0].message.content

# ---------------------------- Render Chat ----------------------------
st.markdown('<div class="chat-container">', unsafe_allow_html=True)
for item in st.session_state.chat_history:
    if isinstance(item, tuple) and len(item) == 3:
        user_msg, ai_msg, audio = item
        st.markdown(f'<div class="chat-bubble-user">{escape(user_msg)}</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="chat-bubble-ai">{escape(ai_msg)}</div>', unsafe_allow_html=True)
        st.markdown(f'''
            <div class="chat-bubble-audio">
            🔊 AI Voice:<br>
            <audio controls src="data:audio/mp3;base64,{audio}"></audio>
            </div>
        ''', unsafe_allow_html=True)
st.markdown('</div>', unsafe_allow_html=True)

# ---------------------------- Fixed Chat Input ----------------------------
st.markdown('<div class="fixed-chat-input">', unsafe_allow_html=True)
with st.form(key="chat_form", clear_on_submit=True):
    chat_input = st.text_area("Your Message", key="chat_input", placeholder="Type your message here...")
    send = st.form_submit_button("Send")
st.markdown('</div>', unsafe_allow_html=True)

if send and chat_input.strip():
    ai_resp = generate_ai_response(chat_input.strip())
    audio_base64 = generate_audio(ai_resp)
    st.session_state.chat_history.append((chat_input.strip(), ai_resp, audio_base64))


# ---------------------------- Word Export ----------------------------
if DOCX_AVAILABLE and st.session_state.chat_history:
    if st.button("📥 Export Chat to Word"):
        doc = Document()
        doc.add_heading("AI Sales Call Assistant Chat Export", 0)
        for user_msg, ai_msg, audio in st.session_state.chat_history:
            doc.add_paragraph(f"User: {user_msg}")
            doc.add_paragraph(f"AI: {ai_msg}\n")
        tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp_file.name)
        with open(tmp_file.name, "rb") as f:
            bytes_data = f.read()
            b64 = base64.b64encode(bytes_data).decode()
            href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="AI_Chat.docx">Click here to download Word file</a>'
            st.markdown(href, unsafe_allow_html=True)
